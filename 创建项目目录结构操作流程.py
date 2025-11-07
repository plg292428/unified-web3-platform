#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""生成UnifiedWeb3Platform项目目录结构操作流程Word文档"""

import os
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("安装python-docx库...")
    import subprocess
    subprocess.check_call(["pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

def add_code_block(doc, code_text):
    """添加代码块"""
    para = doc.add_paragraph()
    para.style = 'No Spacing'
    run = para.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    run._element.rPr.rFonts.set(qn('w:ascii'), 'Consolas')
    run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Consolas')
    para_format = para.paragraph_format
    para_format.left_indent = Inches(0.5)
    para_format.space_before = Pt(6)
    para_format.space_after = Pt(6)

def create_doc():
    doc = Document()
    
    # 设置中文字体
    doc.styles['Normal'].font.name = '宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    # 标题
    title = doc.add_heading('UnifiedWeb3Platform项目目录结构创建操作流程', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('统一Web3平台项目初始化指南')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('')
    
    # 一、项目概述
    doc.add_heading('一、项目概述', 1)
    doc.add_paragraph('UnifiedWeb3Platform是一个统一的Web3全栈平台，整合了以下项目：')
    doc.add_paragraph('• HFastKit：.NET 8工具库', style='List Bullet')
    doc.add_paragraph('• Nblockchain：TRON区块链SDK', style='List Bullet')
    doc.add_paragraph('• SmallTarget：Vue 3企业应用框架', style='List Bullet')
    doc.add_paragraph('• PolygonDapp：Polygon区块链DApp', style='List Bullet')
    
    doc.add_paragraph('')
    doc.add_paragraph('项目目录结构设计：')
    doc.add_paragraph('• src/Frontend/web-app：前端Web应用', style='List Bullet')
    doc.add_paragraph('• src/Backend/UnifiedPlatform.WebApi：Web API服务', style='List Bullet')
    doc.add_paragraph('• src/Backend/UnifiedPlatform.DbService：数据库服务', style='List Bullet')
    doc.add_paragraph('• src/Backend/UnifiedPlatform.Shared：共享库', style='List Bullet')
    doc.add_paragraph('• src/Libraries/HFastKit：HFastKit库', style='List Bullet')
    doc.add_paragraph('• src/Libraries/Nblockchain：Nblockchain库', style='List Bullet')
    
    doc.add_paragraph('')
    
    # 二、环境准备
    doc.add_heading('二、环境准备', 1)
    
    doc.add_heading('2.1 系统要求', 2)
    doc.add_paragraph('• 操作系统：Windows 10/11 或 Windows Server 2016+', style='List Bullet')
    doc.add_paragraph('• PowerShell：5.1+（Windows自带）', style='List Bullet')
    doc.add_paragraph('• 磁盘空间：至少1GB可用空间', style='List Bullet')
    
    doc.add_heading('2.2 验证PowerShell', 2)
    doc.add_paragraph('打开PowerShell，执行以下命令验证：')
    add_code_block(doc, 'powershell -Command "$PSVersionTable.PSVersion"')
    doc.add_paragraph('应显示PowerShell版本号（如：5.1.xxxxx.x）')
    
    doc.add_paragraph('')
    
    # 三、创建项目目录结构
    doc.add_heading('三、创建项目目录结构', 1)
    
    doc.add_heading('3.1 步骤1：创建主项目目录', 2)
    doc.add_paragraph('打开PowerShell，导航到项目根目录（例如：D:\\claude code\\plg），执行：')
    add_code_block(doc, '# 创建主项目目录\nmkdir UnifiedWeb3Platform\n\n# 进入项目目录\ncd UnifiedWeb3Platform')
    
    doc.add_paragraph('')
    doc.add_paragraph('说明：')
    doc.add_paragraph('• mkdir：创建目录命令', style='List Bullet')
    doc.add_paragraph('• 如果目录已存在，PowerShell会提示错误，但可以忽略或使用 -Force 参数', style='List Bullet')
    doc.add_paragraph('• cd：切换目录命令', style='List Bullet')
    
    doc.add_paragraph('')
    
    doc.add_heading('3.2 步骤2：创建前端目录', 2)
    doc.add_paragraph('执行以下命令创建前端项目目录：')
    add_code_block(doc, 'mkdir -p src\\Frontend\\web-app')
    
    doc.add_paragraph('')
    doc.add_paragraph('说明：')
    doc.add_paragraph('• -p：如果父目录不存在，自动创建', style='List Bullet')
    doc.add_paragraph('• Windows PowerShell中，路径分隔符使用反斜杠（\\）', style='List Bullet')
    doc.add_paragraph('• 如果 -p 参数不支持，可以使用以下命令：', style='List Bullet')
    add_code_block(doc, 'New-Item -ItemType Directory -Path "src\\Frontend\\web-app" -Force')
    
    doc.add_paragraph('')
    
    doc.add_heading('3.3 步骤3：创建后端目录', 2)
    doc.add_paragraph('执行以下命令创建后端项目目录：')
    add_code_block(doc, '# 创建Web API目录\nmkdir -p src\\Backend\\UnifiedPlatform.WebApi\n\n# 创建数据库服务目录\nmkdir -p src\\Backend\\UnifiedPlatform.DbService\n\n# 创建共享库目录\nmkdir -p src\\Backend\\UnifiedPlatform.Shared')
    
    doc.add_paragraph('')
    doc.add_paragraph('如果 -p 参数不支持，使用以下命令：')
    add_code_block(doc, 'New-Item -ItemType Directory -Path "src\\Backend\\UnifiedPlatform.WebApi" -Force\nNew-Item -ItemType Directory -Path "src\\Backend\\UnifiedPlatform.DbService" -Force\nNew-Item -ItemType Directory -Path "src\\Backend\\UnifiedPlatform.Shared" -Force')
    
    doc.add_paragraph('')
    
    doc.add_heading('3.4 步骤4：创建库目录', 2)
    doc.add_paragraph('执行以下命令创建库目录：')
    add_code_block(doc, '# 创建HFastKit库目录\nmkdir -p src\\Libraries\\HFastKit\n\n# 创建Nblockchain库目录\nmkdir -p src\\Libraries\\Nblockchain')
    
    doc.add_paragraph('')
    doc.add_paragraph('如果 -p 参数不支持，使用以下命令：')
    add_code_block(doc, 'New-Item -ItemType Directory -Path "src\\Libraries\\HFastKit" -Force\nNew-Item -ItemType Directory -Path "src\\Libraries\\Nblockchain" -Force')
    
    doc.add_paragraph('')
    
    # 四、完整命令脚本
    doc.add_heading('四、完整命令脚本（一键执行）', 1)
    
    doc.add_paragraph('可以将所有命令组合成一个PowerShell脚本，一键执行：')
    add_code_block(doc, '''# UnifiedWeb3Platform项目目录创建脚本
# 执行前请确保在正确的目录（例如：D:\\claude code\\plg）

# 创建主项目目录
if (-not (Test-Path "UnifiedWeb3Platform")) {
    New-Item -ItemType Directory -Path "UnifiedWeb3Platform" -Force
    Write-Host "已创建主项目目录" -ForegroundColor Green
} else {
    Write-Host "主项目目录已存在" -ForegroundColor Yellow
}

# 进入项目目录
Set-Location "UnifiedWeb3Platform"

# 创建前端目录
New-Item -ItemType Directory -Path "src\\Frontend\\web-app" -Force | Out-Null
Write-Host "已创建前端目录" -ForegroundColor Green

# 创建后端目录
New-Item -ItemType Directory -Path "src\\Backend\\UnifiedPlatform.WebApi" -Force | Out-Null
New-Item -ItemType Directory -Path "src\\Backend\\UnifiedPlatform.DbService" -Force | Out-Null
New-Item -ItemType Directory -Path "src\\Backend\\UnifiedPlatform.Shared" -Force | Out-Null
Write-Host "已创建后端目录" -ForegroundColor Green

# 创建库目录
New-Item -ItemType Directory -Path "src\\Libraries\\HFastKit" -Force | Out-Null
New-Item -ItemType Directory -Path "src\\Libraries\\Nblockchain" -Force | Out-Null
Write-Host "已创建库目录" -ForegroundColor Green

# 显示目录结构
Write-Host "`n项目目录结构：" -ForegroundColor Cyan
Get-ChildItem -Path . -Recurse -Directory | Select-Object FullName | Format-Table -AutoSize

Write-Host "`n目录创建完成！" -ForegroundColor Green''')
    
    doc.add_paragraph('')
    doc.add_paragraph('使用方法：')
    doc.add_paragraph('1. 将上述脚本保存为 create_structure.ps1', style='List Number')
    doc.add_paragraph('2. 在PowerShell中执行：', style='List Number')
    add_code_block(doc, 'powershell -ExecutionPolicy Bypass -File create_structure.ps1')
    
    doc.add_paragraph('')
    
    # 五、验证目录结构
    doc.add_heading('五、验证目录结构', 1)
    
    doc.add_paragraph('执行以下命令验证目录是否创建成功：')
    add_code_block(doc, '# 查看当前目录结构\ntree /F /A\n\n# 或者使用PowerShell命令\nGet-ChildItem -Path . -Recurse -Directory | Select-Object FullName')
    
    doc.add_paragraph('')
    doc.add_paragraph('预期输出应包含以下目录：')
    doc.add_paragraph('• UnifiedWeb3Platform\\src\\Frontend\\web-app', style='List Bullet')
    doc.add_paragraph('• UnifiedWeb3Platform\\src\\Backend\\UnifiedPlatform.WebApi', style='List Bullet')
    doc.add_paragraph('• UnifiedWeb3Platform\\src\\Backend\\UnifiedPlatform.DbService', style='List Bullet')
    doc.add_paragraph('• UnifiedWeb3Platform\\src\\Backend\\UnifiedPlatform.Shared', style='List Bullet')
    doc.add_paragraph('• UnifiedWeb3Platform\\src\\Libraries\\HFastKit', style='List Bullet')
    doc.add_paragraph('• UnifiedWeb3Platform\\src\\Libraries\\Nblockchain', style='List Bullet')
    
    doc.add_paragraph('')
    
    # 六、目录结构说明
    doc.add_heading('六、目录结构详细说明', 1)
    
    doc.add_heading('6.1 前端目录 (src/Frontend/web-app)', 2)
    doc.add_paragraph('用途：存放Vue 3前端应用（基于SmallTarget项目）')
    doc.add_paragraph('后续操作：', style='List Bullet')
    doc.add_paragraph('• 将SmallTarget.DappFrontEnd项目复制到此目录', style='List Number')
    doc.add_paragraph('• 安装npm依赖：npm install', style='List Number')
    doc.add_paragraph('• 配置Vite构建工具', style='List Number')
    
    doc.add_paragraph('')
    
    doc.add_heading('6.2 后端目录', 2)
    
    doc.add_heading('6.2.1 UnifiedPlatform.WebApi', 3)
    doc.add_paragraph('用途：Web API服务层，提供RESTful API接口')
    doc.add_paragraph('技术栈：.NET 8 + ASP.NET Core')
    
    doc.add_heading('6.2.2 UnifiedPlatform.DbService', 3)
    doc.add_paragraph('用途：数据库服务层，处理数据访问和业务逻辑')
    doc.add_paragraph('技术栈：.NET 8 + Entity Framework Core')
    
    doc.add_heading('6.2.3 UnifiedPlatform.Shared', 3)
    doc.add_paragraph('用途：共享库，包含通用工具类和模型定义')
    doc.add_paragraph('技术栈：.NET 8类库')
    
    doc.add_paragraph('')
    
    doc.add_heading('6.3 库目录', 2)
    
    doc.add_heading('6.3.1 HFastKit', 3)
    doc.add_paragraph('用途：.NET 8工具库，提供统一响应格式、异常处理等')
    doc.add_paragraph('后续操作：将HFastKit项目复制到此目录')
    
    doc.add_heading('6.3.2 Nblockchain', 3)
    doc.add_paragraph('用途：TRON区块链SDK，提供TRON网络交互功能')
    doc.add_paragraph('后续操作：将Nblockchain项目复制到此目录')
    
    doc.add_paragraph('')
    
    # 七、常见问题
    doc.add_heading('七、常见问题与解决方案', 1)
    
    doc.add_heading('7.1 mkdir -p 命令不支持', 2)
    doc.add_paragraph('问题：PowerShell提示 mkdir -p 参数无效')
    doc.add_paragraph('解决方案：使用 New-Item 命令替代：')
    add_code_block(doc, 'New-Item -ItemType Directory -Path "路径\\目录名" -Force')
    
    doc.add_paragraph('')
    
    doc.add_heading('7.2 目录已存在错误', 2)
    doc.add_paragraph('问题：创建目录时提示目录已存在')
    doc.add_paragraph('解决方案：使用 -Force 参数，或先检查目录是否存在：')
    add_code_block(doc, 'if (-not (Test-Path "目录路径")) {\n    New-Item -ItemType Directory -Path "目录路径" -Force\n}')
    
    doc.add_paragraph('')
    
    doc.add_heading('7.3 权限不足', 2)
    doc.add_paragraph('问题：提示权限不足，无法创建目录')
    doc.add_paragraph('解决方案：')
    doc.add_paragraph('• 以管理员身份运行PowerShell', style='List Bullet')
    doc.add_paragraph('• 检查目标目录的写入权限', style='List Bullet')
    doc.add_paragraph('• 确保不在系统保护目录（如Program Files）中创建', style='List Bullet')
    
    doc.add_paragraph('')
    
    # 八、下一步操作
    doc.add_heading('八、下一步操作建议', 1)
    
    doc.add_paragraph('完成目录结构创建后，建议按以下顺序进行：')
    doc.add_paragraph('1. 复制HFastKit项目到 src/Libraries/HFastKit', style='List Number')
    doc.add_paragraph('2. 复制Nblockchain项目到 src/Libraries/Nblockchain', style='List Number')
    doc.add_paragraph('3. 创建.NET解决方案和项目文件', style='List Number')
    doc.add_paragraph('4. 配置项目引用关系', style='List Number')
    doc.add_paragraph('5. 复制SmallTarget前端项目到 src/Frontend/web-app', style='List Number')
    doc.add_paragraph('6. 安装前端依赖（npm install）', style='List Number')
    doc.add_paragraph('7. 配置开发环境（数据库、Redis等）', style='List Number')
    doc.add_paragraph('8. 编写项目文档和README', style='List Number')
    
    doc.add_paragraph('')
    
    # 九、项目结构图
    doc.add_heading('九、完整项目结构图', 1)
    
    structure_text = '''UnifiedWeb3Platform/
├── src/
│   ├── Frontend/
│   │   └── web-app/                    # Vue 3前端应用
│   │       ├── src/
│   │       ├── public/
│   │       ├── package.json
│   │       └── vite.config.ts
│   │
│   ├── Backend/
│   │   ├── UnifiedPlatform.WebApi/     # Web API服务
│   │   ├── UnifiedPlatform.DbService/ # 数据库服务
│   │   └── UnifiedPlatform.Shared/    # 共享库
│   │
│   └── Libraries/
│       ├── HFastKit/                   # .NET 8工具库
│       └── Nblockchain/                # TRON区块链SDK
│
├── tests/                              # 测试项目（后续创建）
├── docs/                               # 项目文档（后续创建）
├── docker/                             # Docker配置（后续创建）
└── scripts/                            # 部署脚本（后续创建）'''
    
    add_code_block(doc, structure_text)
    
    doc.add_paragraph('')
    
    # 结尾
    doc.add_paragraph('')
    doc.add_paragraph('')
    para = doc.add_paragraph()
    para.add_run('文档生成时间：').bold = True
    from datetime import datetime
    para.add_run(datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    doc.add_paragraph('生成工具：python-docx')
    
    # 保存
    output_dir = Path(__file__).parent
    output_dir.mkdir(parents=True, exist_ok=True)
    output_file = output_dir / 'UnifiedWeb3Platform项目目录结构创建操作流程.docx'
    doc.save(str(output_file))
    
    print(f"Word文档已生成: {output_file}")
    return output_file

if __name__ == '__main__':
    try:
        output = create_doc()
        print(f"\n操作流程文档生成成功！")
        print(f"文件位置: {output}")
    except Exception as e:
        print(f"生成失败: {e}")
        import traceback
        traceback.print_exc()

